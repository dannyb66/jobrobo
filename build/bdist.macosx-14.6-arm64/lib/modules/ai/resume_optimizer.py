import os
import json
import re
from typing import List, Tuple
from openai import OpenAI
from pathlib import Path
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import textwrap
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from modules.config_loader import load_runtime_config, save_runtime_config
from modules.ai.openai_key_loader import load_openai_key
# from config.settings import replace_job_title, rewrite_bullets

class ResumeOptimizer:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)

    def extract_keywords(self, job_description: str) -> List[str]:
        """Extract top 15 keywords from job description using OpenAI API."""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an ATS expert that analyzes job descriptions and "
                            "identifies the most important keywords for ATS systems like Greenhouse and Workday. "
                            "Return exactly 15 most important keywords as a JSON object in this format:\n"
                            '{"keywords": ["keyword1", "keyword2", ..., "keyword15"]}'
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            "Analyze this job description and identify the top 15 most important keywords "
                            "that should be included in a resume to pass ATS systems:\n\n" + job_description
                        )
                    }
                ]
            )

            content = response.choices[0].message.content.strip()
            # print(f"\n🔍 Raw GPT response:\n{content}\n")

            # Try parsing directly
            try:
                result = json.loads(content)
            except json.JSONDecodeError:
                # Try extracting JSON substring
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    result = json.loads(match.group())
                else:
                    raise ValueError("No valid JSON found in response.")

            # print(f"✅ Extracted keywords: {result}")
            return result.get('keywords', [])

        except Exception as e:
            print(f"❌ Error extracting keywords: {e}")
            return []
        
    def read_pdf_resume(self, resume_path: str) -> Tuple[str, List[str], dict]:
        """Read PDF resume and extract the entire skills section with formatting information."""
        try:
            with pdfplumber.open(resume_path) as pdf:
                content = ""
                skills_page = 0
                skills_bbox = None
                current_skills = []

                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    content += page_text + "\n"

                    lines = page_text.splitlines()
                    skills_started = False
                    skills_lines = []

                    for i, line in enumerate(lines):
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_lines.append(line)
                            continue
                        if skills_started:
                            # Stop if we hit the next section (all caps heading or known section)
                            if re.match(r"^[A-Z\s]{5,}$", line.strip()):
                                break
                            skills_lines.append(line)

                    if skills_lines:
                        skills_page = page_num
                        skills_text = " ".join(skills_lines)
                        # Extract the part after "Skills:"
                        skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
                        current_skills = [skill.strip() for skill in skills_raw.split(',') if skill.strip()]

                        # Extract skills bbox
                        words = page.extract_words()
                        for word in words:
                            if "skills" in word['text'].lower():
                                skills_bbox = {
                                    'x0': word['x0'],
                                    'x1': word['x1'],
                                    'top': word['top'],
                                    'bottom': word['bottom']
                                }
                                break

                        break  # We found the section, no need to check other pages

                if not current_skills:
                    raise ValueError("Skills section not found in resume")

                format_info = {
                    'skills_page': skills_page,
                    'skills_bbox': skills_bbox
                }

                # print(f"Skills section found on page {skills_page + 1} with bbox: {skills_bbox}")
                # print(f"Extracted skills: {current_skills}")
                return content, current_skills, format_info

        except Exception as e:
            print(f"Error reading PDF resume: {e}")
            return "", [], {}
        
    def optimize_skills(self, current_skills: List[str], new_keywords: List[str], job_description: str) -> List[str]:
        """Optimize skills list by incorporating new keywords while maintaining the same length."""
        current_set = set(current_skills)
        new_set = set(new_keywords)
        # print(f"Current skills: {current_set}")
        # print(f"New keywords: {new_set}")
        
        # Keywords to add (that aren't already present)
        to_add = new_set - current_set
        # print(f"New keywords to add: {to_add}")
        
        if not to_add:
            return current_skills
            
        # Calculate how many skills we need to remove
        num_to_remove = len(to_add)
        
        # Remove least relevant skills (those not in new_keywords)
        skills_to_keep = list(current_set - to_add)[:len(current_skills) - num_to_remove]
        
        # Combine kept skills with new keywords
        optimized_skills = skills_to_keep + list(to_add)

        # Order the optimized skills with the best order for ATS systems optimized for the job description using OpenAI API
        # try:
        #     response = self.client.chat.completions.create(
        #     model="gpt-4o",
        #     messages=[
        #         {
        #         "role": "system",
        #         "content": (
        #             "You are an ATS optimization expert. Given a list of skills and a job description, "
        #             "reorder the skills to maximize relevance and impact for ATS systems. Make sure to always include all the skills from the original list. "
        #             "Return the reordered skills as a JSON object in this format without the json code block:\n"
        #             '{"ordered_skills": ["skill1", "skill2", ..., "skillN"]}'
        #         )
        #         },
        #         {
        #         "role": "user",
        #         "content": (
        #             f"Here is the job description:\n{job_description}\n\n"
        #             f"Here is the list of skills to reorder:\n{optimized_skills}"
        #         )
        #         }
        #     ]
        #     )

        #     content = response.choices[0].message.content.strip()
        #     print(f"\n🔍 Raw GPT response:\n{content}\n")
        #     # Try parsing the response
        #     try:
        #         result = json.loads(content)
        #         ordered_skills = result.get('ordered_skills', optimized_skills)
        #         if ordered_skills and isinstance(ordered_skills, list):
        #             optimized_skills = ordered_skills
        #     except json.JSONDecodeError:
        #         print("❌ Error parsing OpenAI response for skill ordering. Using original order.")

        # except Exception as e:
        #     print(f"❌ Error reordering skills using OpenAI API: {e}")
        
        return optimized_skills[:len(current_skills)]
    
    def save_optimized_pdf(self, resume_path: str, optimized_skills: List[str], format_info: dict, output_path: str):
        try:
            # Step 1: Use pdfplumber to extract structured text
            with pdfplumber.open(resume_path) as pdf:
                full_text_lines = []
                skills_section_found = False
                skills_started = False

                for page in pdf.pages:
                    page_lines = page.extract_text().splitlines()
                    i = 0
                    while i < len(page_lines):
                        line = page_lines[i]
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_section_found = True

                            # Format the new skills section
                            wrapped_skills = textwrap.wrap(", ".join(optimized_skills), width=110)
                            wrapped_skills[0] = "Skills: " + wrapped_skills[0]  # Add label to first line
                            full_text_lines.extend(wrapped_skills)

                            # Skip original skills block
                            i += 1
                            while i < len(page_lines):
                                next_line = page_lines[i]
                                if re.match(r"^[A-Z\s]{5,}$", next_line.strip()):
                                    break
                                i += 1
                            continue
                        if skills_started and re.match(r"^[A-Z\s]{5,}$", line.strip()):
                            skills_started = False
                        full_text_lines.append(line)
                        i += 1

                if not skills_section_found:
                    raise ValueError("Skills section not found in the resume.")

            # Step 2: Rebuild PDF using ReportLab
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                    rightMargin=40, leftMargin=40,
                                    topMargin=60, bottomMargin=40)
            styles = getSampleStyleSheet()
            story = []

            for para in "\n".join(full_text_lines).split('\n\n'):
                for line in para.strip().splitlines():
                    story.append(Paragraph(line.strip(), styles["Normal"]))

            doc.build(story)

            # Step 3: Save to disk
            with open(output_path, 'wb') as f:
                f.write(buffer.getvalue())

            print(f"✅ Successfully saved updated PDF to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized PDF resume: {e}")

    def read_docx_resume(self, resume_path: str) -> Tuple[Document, List[str]]:
        """Read a DOCX resume and extract the skills section."""
        try:
            doc = Document(resume_path)
            skills_text = ""
            skills_started = False
            skills_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not skills_started and re.match(r"(?i)^skills\s*:", text):
                    skills_started = True
                    skills_lines.append(text)
                    continue
                if skills_started:
                    if re.match(r"^[A-Z\s]{5,}$", text):
                        break
                    skills_lines.append(text)

            if not skills_lines:
                raise ValueError("Skills section not found in DOCX resume")

            skills_text = " ".join(skills_lines)
            skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
            current_skills = [s.strip() for s in skills_raw.split(',') if s.strip()]

            return doc, current_skills

        except Exception as e:
            print(f"Error reading DOCX resume: {e}")
            return None, []
        
    # Replace job titles in the Professional Experience section
    def replace_job_titles(self, doc: Document, new_title: str, max_line_length: int = 150) -> None:
        """
        Replace all job titles in the 'Professional Experience' section with the given new_title,
        padding it with spaces so the final line reaches up to max_line_length without overflowing.
        """
        try:
            is_in_experience_section = False
            job_title_with_dates_pattern = re.compile(r'^(.*?)(\s{2,}|\t+)?\(?[A-Za-z]+\s+\d{4}.*\)?$')

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Toggle section flags
                if text.upper() == "PROFESSIONAL EXPERIENCE":
                    is_in_experience_section = True
                    continue
                elif text.upper() in ["EDUCATION", "ENTREPRENEURIAL VENTURES", "ACADEMIC PROJECTS", "CERTIFICATIONS", "SKILLS"]:
                    is_in_experience_section = False
                    continue

                if is_in_experience_section:
                    match = job_title_with_dates_pattern.match(text)
                    if match:
                        old_title = match.group(1).strip()
                        rest_of_line = text[len(old_title):].strip()

                        # Limit new title to fit within line length
                        allowed_title_length = max_line_length - len(rest_of_line)
                        safe_new_title = new_title[:allowed_title_length].rstrip()

                        # Clean up the title, no extra spaces, no special characters or numbers
                        safe_new_title = re.sub(r'[^A-Za-z\s]', '', safe_new_title).strip()

                        # Calculate how many spaces to add after title
                        padding_space = max(1, max_line_length - len(safe_new_title + ' ' + rest_of_line))
                        padded_title = safe_new_title + ' ' * padding_space                        

                        # Update paragraph text with a tab between the title and rest of the line, and apply italic style
                        # para.text = padded_title + rest_of_line
                        para.text = safe_new_title + "\t" + rest_of_line.lstrip()
                        run = para.runs[0]
                        run.italic = True
                        run.font.size = Pt(9)

        except Exception as e:
            print(f"❌ Error replacing job titles: {e}")

    # Rewrite experience bullets to match job description
    def rewrite_experience_bullets(self, doc: Document, job_title: str, job_description: str, optimized_skills: List[str]) -> None:
        """Rewrite experience bullets to match job description using OpenAI API."""
        try:
            experience_bullets = self._extract_experience_bullets(doc)
            print(f"🔍 Extracted experience bullets: {experience_bullets}")
            if not experience_bullets:
                raise ValueError("No experience bullets found in the resume.")

            rewritten_experience = self._call_openai_to_rewrite_bullets(
                experience_bullets, job_title, job_description
            )
            print(f"🔍 Rewritten experience bullets: {rewritten_experience}")

            if rewritten_experience:
                self._update_document_with_rewritten_bullets(doc, experience_bullets, rewritten_experience)

        except Exception as e:
            print(f"❌ Error rewriting experience bullets: {e}")

    def _extract_experience_bullets(self, doc: Document) -> List[dict]:
        """Extract experience bullets from the document."""
        experience_bullets = []
        current_entry = self._initialize_experience_entry()
        is_in_experience_section = False
        header_line_stage = 0  # 0 = expecting company, 1 = expecting title, 2 = collecting bullets

        # print("\n📄 Starting experience bullet extraction...")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # print(f"\n🔎 Analyzing line: '{text}'")

            if self._is_section_header(text, "PROFESSIONAL EXPERIENCE"):
                # print("📌 Entered PROFESSIONAL EXPERIENCE section.")
                is_in_experience_section = True
                header_line_stage = 0
                continue
            elif self._is_section_header(text, ["EDUCATION", "ENTREPRENEURIAL VENTURES", "ACADEMIC PROJECTS", "CERTIFICATIONS", "SKILLS"]):
                # if is_in_experience_section:
                #     print("📤 Exiting PROFESSIONAL EXPERIENCE section.")
                is_in_experience_section = False
                self._finalize_experience_entry(current_entry, experience_bullets)
                current_entry = self._initialize_experience_entry()
                continue

            if is_in_experience_section:
                if self._is_header_line(text):
                    # print(f"🧾 Detected header line: '{text}'")
                    if header_line_stage == 0:
                        self._finalize_experience_entry(current_entry, experience_bullets)
                        current_entry = self._initialize_experience_entry()
                        current_entry["company"] = text.split("\t")[0].strip()
                        header_line_stage = 1
                    elif header_line_stage == 1:
                        current_entry["job_title"] = text.split("\t")[0].strip()
                        header_line_stage = 2
                    elif header_line_stage == 2:
                        # New workex block started – reset everything
                        # print("🔁 Starting new work experience entry.")
                        self._finalize_experience_entry(current_entry, experience_bullets)
                        current_entry = self._initialize_experience_entry()
                        current_entry["company"] = text.split("\t")[0].strip()
                        header_line_stage = 1  # Expect job title next
                else:
                    if header_line_stage == 2:
                        # print(f"➕ Adding bullet: '{text}'")
                        self._add_bullet_to_entry(text, current_entry)
                    else:
                        print(f"⚠️ Skipping unexpected line: '{text}'")

        # print("\n🧹 Finalizing last experience entry if any.")
        self._finalize_experience_entry(current_entry, experience_bullets)

        # print(f"\n✅ Extracted {len(experience_bullets)} experience entries.")
        for idx, entry in enumerate(experience_bullets):
            print(f"  {idx+1}. Company: {entry['company']}, Title: {entry['job_title']}, Bullets: {len(entry['bullets'])}")

        return experience_bullets

    def _initialize_experience_entry(self) -> dict:
        """Initialize a new experience entry."""
        return {"company": None, "job_title": None, "bullets": []}

    def _finalize_experience_entry(self, entry: dict, experience_bullets: List[dict]) -> None:
        """Finalize the current experience entry and add it to the list."""
        if entry["company"] and entry["job_title"] and entry["bullets"]:
            experience_bullets.append(entry)

    def _is_section_header(self, text: str, headers: List[str] or str) -> bool:
        """Check if the text matches a section header."""
        if isinstance(headers, list):
            return text.upper() in headers
        return text.upper() == headers

    def _is_header_line(self, text: str) -> bool:
        """Determine if the line is a header (company or job title)."""
        return "\t" in text or re.search(r"\s{2,}", text)

    def _process_header_line(self, text: str, current_entry: dict) -> dict:
        """Process a header line and update the current entry."""
        if current_entry["company"] is None:
            current_entry["company"] = text.split("\t")[0].strip()
            current_entry["job_title"] = None
            current_entry["bullets"] = []
        else:
            current_entry["job_title"] = text.split("\t")[0].strip()
            current_entry["bullets"] = []
        return current_entry

    def _add_bullet_to_entry(self, text: str, current_entry: dict) -> None:
        """Add a bullet point to the current entry."""
        current_entry["bullets"].append(text)

    def _call_openai_to_rewrite_bullets(self, experience_bullets: List[dict], job_title: str, job_description: str) -> List[dict]:
        """Call OpenAI API to rewrite experience bullets."""
        try:
            system_message = self._generate_system_message()
            user_message = self._generate_user_message(experience_bullets, job_title, job_description)

            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ]
            )

            return self._parse_openai_response(response, experience_bullets)

        except Exception as e:
            print(f"❌ Error calling OpenAI API to rewrite bullets: {e}")
            return experience_bullets

    def _generate_system_message(self) -> str:
        """Generate the system message for the OpenAI API request."""
        return (
            "You are an ATS resume optimization expert and professional career writer. "
            "You are given a list of professional experience entries, where each entry contains:\n"
            "- the job title\n"
            "- the company name\n"
            "- a list of bullet points describing the work.\n\n"
            "Your task is to rewrite each bullet point **specifically for a target job title and description**, "
            "while making sure the rewritten content still logically fits the role and domain of the original company and job title.\n\n"
            "Use keywords, technologies, and phrasing relevant to the target job description wherever appropriate, "
            "without introducing false or irrelevant information. If a bullet is about a different domain, find the closest possible conceptual match "
            "to bridge relevance. Optimize each bullet to align better with the responsibilities and qualifications expected in the target job.\n\n"
            "Bullet Rewriting Guidelines:\n"
            "1. **Use either of these styles:**\n"
            "   - [A] Action + Project + Result\n"
            "     → Example: 'Optimized backend APIs for a fintech product, reducing latency by 30% and improving customer satisfaction scores.'\n"
            "   - [B] Accomplished [X] as measured by [Y] by doing [Z]\n"
            "     → Example: 'Increased automation coverage by 40% by designing and deploying scalable AI workflows.'\n\n"
            "2. **Do NOT exceed the original bullet's character count**.\n"
            "3. **Use a minimum of 120 characters per bullet.**\n"
            "4. Ensure each bullet reflects **individual contributions**, not team achievements.\n"
            "5. Use **concise, powerful, and ATS-friendly phrasing**.\n"
            "6. Keep the job title as the one in the orignal bullets and do not update it to the new job title.\n"
            "7. Strongly align with the **job title and job description**.\n\n"
            "Return only a valid JSON object in this format:\n"
            "{ \"rewritten_experience\": [ { \"job_title\": \"...\", \"company\": \"...\", \"bullets\": [\"...\", \"...\"] }, ... ] }"
        )

    def _generate_user_message(self, experience_bullets: List[dict], job_title: str, job_description: str) -> str:
        """Generate the user message for the OpenAI API request."""
        return (
            f"Here is the job title: {job_title}\n\n"
            f"Here is the job description:\n{job_description}\n\n"
            f"Here are the experience entries to rewrite:\n{json.dumps(experience_bullets, indent=2)}"
        )

    def _parse_openai_response(self, response, fallback: List[dict]) -> List[dict]:
        """Parse the OpenAI API response and extract rewritten experience bullets."""
        try:
            content = response.choices[0].message.content.strip()

            # Remove JSON code block markers if present
            if content.startswith("```json") and content.endswith("```"):
                content = content[7:-3].strip()

            # Parse the response
            result = json.loads(content)
            return result.get('rewritten_experience', fallback)

        except (json.JSONDecodeError, AttributeError) as e:
            print(f"❌ Error parsing OpenAI API response: {e}")
            return fallback
        
    def _update_document_with_rewritten_bullets(self, doc: Document, original_bullets: List[dict], rewritten_bullets: List[dict]) -> None:
        """Update the document with rewritten bullets."""
        bullet_lookup = self._create_bullet_lookup(original_bullets)
        rewritten_lookup = self._create_bullet_lookup(rewritten_bullets)

        current_company = None
        current_job_title = None
        is_in_target_entry = False
        bullet_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if self._is_company_line(text):
                current_company = self._extract_company_name(text)
                current_job_title = None
                is_in_target_entry = False
                bullet_index = 0

            elif current_company and self._is_job_title_line(text):
                current_job_title = self._extract_job_title(text)
                bullet_key = (current_company, current_job_title)
                is_in_target_entry = bullet_key in bullet_lookup and bullet_key in rewritten_lookup
                bullet_index = 0

            elif is_in_target_entry:
                self._update_bullet_point(
                    para, bullet_lookup, rewritten_lookup, current_company, current_job_title, bullet_index
                )
                bullet_index += 1

                if self._has_updated_all_bullets(bullet_index, rewritten_lookup, current_company, current_job_title):
                    current_company, current_job_title, is_in_target_entry, bullet_index = self._reset_state()

            elif current_company and current_job_title and not is_in_target_entry:
                current_company, current_job_title, is_in_target_entry, bullet_index = self._reset_state()

    def _create_bullet_lookup(self, bullets: List[dict]) -> dict:
        """Create a lookup dictionary for bullets."""
        return {
            (entry['company'], entry['job_title']): entry['bullets']
            for entry in bullets
        }

    def _is_company_line(self, text: str) -> bool:
        """Check if the line represents a company."""
        return re.fullmatch(r"[A-Z0-9\s&/'’*()\-]+", text.split("\t")[0]) and (
            "\t" in text or re.search(r"\s{2,}", text)
        )

    def _extract_company_name(self, text: str) -> str:
        """Extract the company name from the text."""
        return text.split("\t")[0].strip()

    def _is_job_title_line(self, text: str) -> bool:
        """Check if the line represents a job title."""
        return "\t" in text or re.search(r"\s{2,}", text)

    def _extract_job_title(self, text: str) -> str:
        """Extract the job title from the text."""
        return text.split("\t")[0].strip()

    def _update_bullet_point(
        self, para, bullet_lookup, rewritten_lookup, current_company, current_job_title, bullet_index
    ) -> None:
        """Update a single bullet point in the document."""
        old_bullets = bullet_lookup[(current_company, current_job_title)]
        new_bullets = rewritten_lookup[(current_company, current_job_title)]

        if bullet_index < len(old_bullets) and bullet_index < len(new_bullets):
            para.text = new_bullets[bullet_index]
            for run in para.runs:
                run.font.size = Pt(9.5)

    def _has_updated_all_bullets(self, bullet_index: int, rewritten_lookup: dict, current_company: str, current_job_title: str) -> bool:
        """Check if all bullets for the current entry have been updated."""
        return bullet_index >= len(rewritten_lookup[(current_company, current_job_title)])

    def _reset_state(self) -> Tuple[None, None, bool, int]:
        """Reset the state variables."""
        return None, None, False, 0

    """Calculate page count for a given document."""
    def _calculate_page_count(self, doc: Document, output_path: str) -> int:
        temp_path = str(Path(output_path).parent / "temp.docx")
        doc.save(temp_path)
        return len(Document(temp_path).paragraphs)

    def save_optimized_docx(self, doc: Document, optimized_skills: List[str], output_path: str, job_title: str = "job", job_description: str = "") -> None:
        """Update the skills section in the DOCX and save it."""
        try:
            skills_index, end_index = self._find_skills_section(doc)
            if skills_index is None:
                raise ValueError("Skills section not found in DOCX resume")

            self._update_skills_section(doc, skills_index, end_index, optimized_skills)            

            if rewrite_bullets:
                original_page_count = self._calculate_page_count(doc, output_path)
                while True:
                    self.rewrite_experience_bullets(doc, job_title, job_description, optimized_skills)
                    temp_path = str(Path(output_path).parent / "temp.docx")
                    doc.save(temp_path)
                    updated_page_count = self._calculate_page_count(Document(temp_path), output_path)
                    if updated_page_count <= original_page_count:
                        break

            if replace_job_title:
                self.replace_job_titles(doc, job_title)

            doc.save(output_path)
            print(f"✅ Successfully saved updated DOCX resume to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized DOCX resume: {e}")

    def _find_skills_section(self, doc: Document) -> Tuple[int, int]:
        """Find the start and end indices of the skills section."""
        skills_started = False
        skills_index = None
        end_index = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not skills_started and re.match(r"(?i)^skills\s*:", text):
                skills_started = True
                skills_index = i
                continue
            if skills_started and re.match(r"^[A-Z\s]{5,}$", text):
                end_index = i
                break

        if end_index is None:
            end_index = len(doc.paragraphs)

        return skills_index, end_index

    def _update_skills_section(self, doc: Document, skills_index: int, end_index: int, optimized_skills: List[str]) -> None:
        """Update the skills section in the document."""
        # Remove existing skills content
        for i in range(end_index - 1, skills_index, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)

        # Add optimized skills
        skills_text_wrapped = textwrap.wrap(", ".join(optimized_skills), width=120)
        doc.paragraphs[skills_index].text = skills_text_wrapped[0]

        # Add additional lines for wrapped skills
        insert_after = doc.paragraphs[skills_index]._element
        for line in skills_text_wrapped[1:]:
            new_para = doc.add_paragraph(line)
            new_para.style = doc.paragraphs[skills_index].style
            insert_after.addnext(new_para._element)
            insert_after = new_para._element

    # Convert docx to pdf
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str):
        try:
            # pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
            # docx2pdf.convert(docx_path, pdf_path)
            # Use LibreOffice to convert DOCX to PDF
            import subprocess
            libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', str(Path(pdf_path).parent)], check=True)
            print(f"✅ Successfully converted DOCX to PDF: {pdf_path}")
        except ImportError:
            print("❌ pypandoc module not found. Please install it using 'pip install pypandoc'.")
        except Exception as e:
            print(f"❌ Error converting DOCX to PDF: {e}")

# Function to run the resume optimization process given job description and resume path similar to main()
def run_resume_optimization(job_description: str, resume_path: str, job_id: str, job_title: str = "job", first_name: str = "first", last_name: str = "last"):
    # Get OpenAI API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Please set the OPENAI_API_KEY environment variable")
        return

    optimizer = ResumeOptimizer(api_key)

    # Extract keywords from job description
    print("\nExtracting keywords from job description...")
    keywords = optimizer.extract_keywords(job_description)
    print(f"Top keywords identified: {', '.join(keywords)}")

    # Check if resume file exists
    if not Path(resume_path).exists():
        print("Resume file not found!")
        return

    ext = Path(resume_path).suffix.lower()
    output_path = None

    # Process resume based on file type
    print("\nProcessing resume...")
    # Sanitize variables to ensure they don't contain invalid path characters
    sanitized_job_title = re.sub(r'[\/:*?"<>|]', '_', job_title.lower().replace(' ', '_'))
    sanitized_first_name = re.sub(r'[\/:*?"<>|]', '_', first_name.lower())
    sanitized_last_name = re.sub(r'[\/:*?"<>|]', '_', last_name.lower())

    if ext == ".pdf":
        content, current_skills, format_info = optimizer.read_pdf_resume(resume_path)
        if not content or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.pdf"
        optimizer.save_optimized_pdf(resume_path, optimized_skills, format_info, str(output_path))

    elif ext == ".docx":
        doc, current_skills = optimizer.read_docx_resume(resume_path)
        if not doc or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path_docx = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.docx"
        output_path_pdf = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.pdf"
        optimizer.save_optimized_docx(doc, optimized_skills, str(output_path_docx), job_title, job_description)
        optimizer.convert_docx_to_pdf(str(output_path_docx), str(output_path_pdf))
        output_path = output_path_pdf

    else:
        print("❌ Unsupported file format. Please use PDF or DOCX.")
        return

    print(f"\n🎯 Optimized resume saved to: {output_path}")
    return output_path

# Function to load configuration variables
def load_config() -> dict:
    """Load configuration variables from a JSON file."""
    config_path = Path(__file__).parent.parent.parent / "config" / "resume_optimizer_defaults.json"
    try:
        with open(config_path, "r") as config_file:
            config = json.load(config_file)
            return {
                "replace_job_title": config.get("replace_job_title", False),
                "rewrite_bullets": config.get("rewrite_bullets", False),
                "resume_path": config.get("resume_path", ""),
                "job_id": config.get("job_id", ""),
                "first_name": config.get("first_name", "first"),
                "last_name": config.get("last_name", "last"),
                "job_title": config.get("job_title", "job"),
                "job_description": config.get("job_description", ""),
            }
    except FileNotFoundError:
        print(f"❌ Config file not found at {config_path}. Using default values.")
        return {
            "replace_job_title": False,
            "rewrite_bullets": False,
            "resume_path": "",
            "job_id": "",
            "first_name": "first",
            "last_name": "last",
            "job_title": "job",
            "job_description": "",
        }
    except json.JSONDecodeError:
        print(f"❌ Error decoding JSON config file at {config_path}. Using default values.")
        return {
            "replace_job_title": False,
            "rewrite_bullets": False,
            "resume_path": "",
            "job_id": "",
            "first_name": "first",
            "last_name": "last",
            "job_title": "job",
            "job_description": "",
        }

# Load configuration variables
# load_dotenv()  # Load environment variables from .env file
load_openai_key()  # Load OpenAI API key from cloud or .env file
# config = load_config()
config = load_runtime_config()
replace_job_title = config["replace_job_title"]
rewrite_bullets = config["rewrite_bullets"]
job_description = config["job_description"]
resume_path = config["resume_path"]
job_id = config["job_id"]
job_title = config["job_title"]
first_name = config["first_name"]
last_name = config["last_name"]

# Main function to run the resume optimization process
if __name__ == "__main__":
    # Run the resume optimization process
    run_resume_optimization(job_description, resume_path, job_id, job_title, first_name, last_name)
