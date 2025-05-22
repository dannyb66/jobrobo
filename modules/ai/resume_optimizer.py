import os
import json
import re
from typing import List, Tuple
from openai import OpenAI
from pathlib import Path
import pdfplumber
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import textwrap
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt
from config.settings import replace_job_title, rewrite_bullets

class ResumeOptimizer:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)

    def extract_keywords(self, job_description: str) -> List[str]:
        """Extract top 15 keywords from job description using OpenAI API."""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an ATS expert that analyzes job descriptions and "
                            "identifies the most important keywords for ATS systems like Greenhouse and Workday. "
                            "Return exactly 15 most important keywords as a JSON object in this format:\n"
                            '{"keywords": ["keyword1", "keyword2", ..., "keyword15"]}'
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            "Analyze this job description and identify the top 15 most important keywords "
                            "that should be included in a resume to pass ATS systems:\n\n" + job_description
                        )
                    }
                ]
            )

            content = response.choices[0].message.content.strip()
            # print(f"\n🔍 Raw GPT response:\n{content}\n")

            # Try parsing directly
            try:
                result = json.loads(content)
            except json.JSONDecodeError:
                # Try extracting JSON substring
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    result = json.loads(match.group())
                else:
                    raise ValueError("No valid JSON found in response.")

            # print(f"✅ Extracted keywords: {result}")
            return result.get('keywords', [])

        except Exception as e:
            print(f"❌ Error extracting keywords: {e}")
            return []
        
    def read_pdf_resume(self, resume_path: str) -> Tuple[str, List[str], dict]:
        """Read PDF resume and extract the entire skills section with formatting information."""
        try:
            with pdfplumber.open(resume_path) as pdf:
                content = ""
                skills_page = 0
                skills_bbox = None
                current_skills = []

                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    content += page_text + "\n"

                    lines = page_text.splitlines()
                    skills_started = False
                    skills_lines = []

                    for i, line in enumerate(lines):
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_lines.append(line)
                            continue
                        if skills_started:
                            # Stop if we hit the next section (all caps heading or known section)
                            if re.match(r"^[A-Z\s]{5,}$", line.strip()):
                                break
                            skills_lines.append(line)

                    if skills_lines:
                        skills_page = page_num
                        skills_text = " ".join(skills_lines)
                        # Extract the part after "Skills:"
                        skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
                        current_skills = [skill.strip() for skill in skills_raw.split(',') if skill.strip()]

                        # Extract skills bbox
                        words = page.extract_words()
                        for word in words:
                            if "skills" in word['text'].lower():
                                skills_bbox = {
                                    'x0': word['x0'],
                                    'x1': word['x1'],
                                    'top': word['top'],
                                    'bottom': word['bottom']
                                }
                                break

                        break  # We found the section, no need to check other pages

                if not current_skills:
                    raise ValueError("Skills section not found in resume")

                format_info = {
                    'skills_page': skills_page,
                    'skills_bbox': skills_bbox
                }

                # print(f"Skills section found on page {skills_page + 1} with bbox: {skills_bbox}")
                # print(f"Extracted skills: {current_skills}")
                return content, current_skills, format_info

        except Exception as e:
            print(f"Error reading PDF resume: {e}")
            return "", [], {}
        
    def optimize_skills(self, current_skills: List[str], new_keywords: List[str], job_description: str) -> List[str]:
        """Optimize skills list by incorporating new keywords while maintaining the same length."""
        current_set = set(current_skills)
        new_set = set(new_keywords)
        # print(f"Current skills: {current_set}")
        # print(f"New keywords: {new_set}")
        
        # Keywords to add (that aren't already present)
        to_add = new_set - current_set
        # print(f"New keywords to add: {to_add}")
        
        if not to_add:
            return current_skills
            
        # Calculate how many skills we need to remove
        num_to_remove = len(to_add)
        
        # Remove least relevant skills (those not in new_keywords)
        skills_to_keep = list(current_set - to_add)[:len(current_skills) - num_to_remove]
        
        # Combine kept skills with new keywords
        optimized_skills = skills_to_keep + list(to_add)

        # Order the optimized skills with the best order for ATS systems optimized for the job description using OpenAI API
        # try:
        #     response = self.client.chat.completions.create(
        #     model="gpt-4o",
        #     messages=[
        #         {
        #         "role": "system",
        #         "content": (
        #             "You are an ATS optimization expert. Given a list of skills and a job description, "
        #             "reorder the skills to maximize relevance and impact for ATS systems. Make sure to always include all the skills from the original list. "
        #             "Return the reordered skills as a JSON object in this format without the json code block:\n"
        #             '{"ordered_skills": ["skill1", "skill2", ..., "skillN"]}'
        #         )
        #         },
        #         {
        #         "role": "user",
        #         "content": (
        #             f"Here is the job description:\n{job_description}\n\n"
        #             f"Here is the list of skills to reorder:\n{optimized_skills}"
        #         )
        #         }
        #     ]
        #     )

        #     content = response.choices[0].message.content.strip()
        #     print(f"\n🔍 Raw GPT response:\n{content}\n")
        #     # Try parsing the response
        #     try:
        #         result = json.loads(content)
        #         ordered_skills = result.get('ordered_skills', optimized_skills)
        #         if ordered_skills and isinstance(ordered_skills, list):
        #             optimized_skills = ordered_skills
        #     except json.JSONDecodeError:
        #         print("❌ Error parsing OpenAI response for skill ordering. Using original order.")

        # except Exception as e:
        #     print(f"❌ Error reordering skills using OpenAI API: {e}")
        
        return optimized_skills[:len(current_skills)]
    
    def save_optimized_pdf(self, resume_path: str, optimized_skills: List[str], format_info: dict, output_path: str):
        try:
            # Step 1: Use pdfplumber to extract structured text
            with pdfplumber.open(resume_path) as pdf:
                full_text_lines = []
                skills_section_found = False
                skills_started = False

                for page in pdf.pages:
                    page_lines = page.extract_text().splitlines()
                    i = 0
                    while i < len(page_lines):
                        line = page_lines[i]
                        if not skills_started and re.match(r"(?i)^skills\s*:", line.strip()):
                            skills_started = True
                            skills_section_found = True

                            # Format the new skills section
                            wrapped_skills = textwrap.wrap(", ".join(optimized_skills), width=110)
                            wrapped_skills[0] = "Skills: " + wrapped_skills[0]  # Add label to first line
                            full_text_lines.extend(wrapped_skills)

                            # Skip original skills block
                            i += 1
                            while i < len(page_lines):
                                next_line = page_lines[i]
                                if re.match(r"^[A-Z\s]{5,}$", next_line.strip()):
                                    break
                                i += 1
                            continue
                        if skills_started and re.match(r"^[A-Z\s]{5,}$", line.strip()):
                            skills_started = False
                        full_text_lines.append(line)
                        i += 1

                if not skills_section_found:
                    raise ValueError("Skills section not found in the resume.")

            # Step 2: Rebuild PDF using ReportLab
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                    rightMargin=40, leftMargin=40,
                                    topMargin=60, bottomMargin=40)
            styles = getSampleStyleSheet()
            story = []

            for para in "\n".join(full_text_lines).split('\n\n'):
                for line in para.strip().splitlines():
                    story.append(Paragraph(line.strip(), styles["Normal"]))

            doc.build(story)

            # Step 3: Save to disk
            with open(output_path, 'wb') as f:
                f.write(buffer.getvalue())

            print(f"✅ Successfully saved updated PDF to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized PDF resume: {e}")

    def read_docx_resume(self, resume_path: str) -> Tuple[Document, List[str]]:
        """Read a DOCX resume and extract the skills section."""
        try:
            doc = Document(resume_path)
            skills_text = ""
            skills_started = False
            skills_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not skills_started and re.match(r"(?i)^skills\s*:", text):
                    skills_started = True
                    skills_lines.append(text)
                    continue
                if skills_started:
                    if re.match(r"^[A-Z\s]{5,}$", text):
                        break
                    skills_lines.append(text)

            if not skills_lines:
                raise ValueError("Skills section not found in DOCX resume")

            skills_text = " ".join(skills_lines)
            skills_raw = re.sub(r"(?i)^skills\s*:\s*", "", skills_text)
            current_skills = [s.strip() for s in skills_raw.split(',') if s.strip()]

            return doc, current_skills

        except Exception as e:
            print(f"Error reading DOCX resume: {e}")
            return None, []
        
    # Replace job titles in the Professional Experience section
    def replace_job_titles(self, doc: Document, new_title: str, max_line_length: int = 150) -> None:
        """
        Replace all job titles in the 'Professional Experience' section with the given new_title,
        padding it with spaces so the final line reaches up to max_line_length without overflowing.
        """
        try:
            is_in_experience_section = False
            job_title_with_dates_pattern = re.compile(r'^(.*?)(\s{2,}|\t+)?\(?[A-Za-z]+\s+\d{4}.*\)?$')

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Toggle section flags
                if text.upper() == "PROFESSIONAL EXPERIENCE":
                    is_in_experience_section = True
                    continue
                elif text.upper() in ["EDUCATION", "ENTREPRENEURIAL VENTURES", "ACADEMIC PROJECTS", "CERTIFICATIONS", "SKILLS"]:
                    is_in_experience_section = False
                    continue

                if is_in_experience_section:
                    match = job_title_with_dates_pattern.match(text)
                    if match:
                        old_title = match.group(1).strip()
                        rest_of_line = text[len(old_title):].strip()

                        # Limit new title to fit within line length
                        allowed_title_length = max_line_length - len(rest_of_line)
                        safe_new_title = new_title[:allowed_title_length].rstrip()

                        # Clean up the title, no extra spaces, no special characters or numbers
                        safe_new_title = re.sub(r'[^A-Za-z\s]', '', safe_new_title).strip()

                        # Calculate how many spaces to add after title
                        padding_space = max(1, max_line_length - len(safe_new_title + ' ' + rest_of_line))
                        padded_title = safe_new_title + ' ' * padding_space                        

                        # Update paragraph text with a tab between the title and rest of the line, and apply italic style
                        # para.text = padded_title + rest_of_line
                        para.text = safe_new_title + "\t" + rest_of_line.lstrip()
                        run = para.runs[0]
                        run.italic = True
                        run.font.size = Pt(9)

        except Exception as e:
            print(f"❌ Error replacing job titles: {e}")

    # Rewrite experience bullets to match job description
    def rewrite_experience_bullets(self, doc: Document, job_title: str, job_description: str, optimized_skills: List[str]) -> None:
        """Rewrite experience bullets to match job description using OpenAI API."""
        try:
            is_in_experience_section = False
            experience_bullets = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Toggle section flags
                if text.upper() == "PROFESSIONAL EXPERIENCE":
                    is_in_experience_section = True
                    continue
                elif text.upper() in ["EDUCATION", "ENTREPRENEURIAL VENTURES", "ACADEMIC PROJECTS", "CERTIFICATIONS", "SKILLS"]:
                    is_in_experience_section = False
                    continue

                if is_in_experience_section and text and "\t" not in text:
                    experience_bullets.append(text)

            # print(f"Extracted experience bullets: {experience_bullets}")
            if not experience_bullets:
                raise ValueError("No experience bullets found in the resume.")

            # Call OpenAI API to rewrite bullets

            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an ATS resume optimization expert and professional career writer. "
                            "Given a list of experience bullets, your task is to rewrite each one to better match a specific job description, "
                            "while staying within the original length of each bullet and making it a great fit for the job description. "
                            "Make the language more aligned with the target role by incorporating relevant terminology, technologies, and metrics. "
                            "Feel free to write it with a more job description-specific and job title-specific bullet while it is related to the company's work the bullet reference to."                            
                            "\n\n"
                            "Follow these bullet rewriting guidelines:\n"
                            "1. Use the Action + Project + Result format:\n"
                            "   - [A] Start with a strong action verb.\n"
                            "   - [P] Clearly describe the project or task.\n"
                            "   - [R] State the results or impact, using quantifiable metrics where possible.\n"
                            "   Example: 'Optimized backend APIs for a fintech product, reducing latency by 30% and improving customer satisfaction scores.'\n\n"
                            "2. Alternatively, use the Accomplished [X] as measured by [Y] by doing [Z] format:\n"
                            "   - [X] Start with the impact or accomplishment.\n"
                            "   - [Y] Quantify the improvement (%, time saved, dollars, etc.).\n"
                            "   - [Z] Describe the specific contribution or action.\n"
                            "   Example: 'Increased automation coverage by 40% by designing and deploying scalable AI workflows.'\n\n"
                            "3. DO NOT exceed the character length of the original bullet.\n"
                            "4. Minimum of 120 characters should be used.\n"
                            "Ensure alignment with keywords and skills mentioned in the job description (e.g., quantitative analysis, portfolio risk, Python, SaaS, financial modeling).\n"
                            "Each bullet should reflect **individual contributions**, not team achievements.\n"
                            "Use concise, powerful phrasing optimized for ATS scanning.\n\n"
                            "Return the output as a JSON object in this format:\n"
                            "{ \"rewritten_bullets\": [\"bullet1\", \"bullet2\", ..., \"bulletN\"] }"
                        )
                      },
                    {
                        "role": "user",
                        "content": (
                            f"Here is the job title: {job_title}\n\n"
                            f"Here is the job description:\n{job_description}\n\n"
                            f"Here are the experience bullets to rewrite:\n{experience_bullets}"
                        )
                    }
                ]
            )

            content = response.choices[0].message.content.strip()
            # print(f"\n🔍 Raw GPT response:\n{content}\n")

            # Remove JSON code block markers if present
            if content.startswith("```json") and content.endswith("```"):
                content = content[7:-3].strip()

            # Try parsing directly
            try:
                result = json.loads(content)
                rewritten_bullets = result.get('rewritten_bullets', experience_bullets)
                if rewritten_bullets and isinstance(rewritten_bullets, list):
                    for i, para in enumerate(doc.paragraphs):
                        if para.text.strip() in experience_bullets:
                            para.text = rewritten_bullets[experience_bullets.index(para.text.strip())]
                            for run in para.runs:
                                run.font.size = Pt(9.5)  # Set the font size to match the previous text
            except json.JSONDecodeError:
                raise ValueError("No valid JSON found in response.")
            # print(f"✅ Rewritten bullets: {rewritten_bullets}")
        except Exception as e:
            print(f"❌ Error rewriting experience bullets: {e}")
            return
        # If the API call fails, keep the original bullets
        # for i, para in enumerate(doc.paragraphs):
        #     if para.text.strip() in experience_bullets:
        #         para.text = experience_bullets[experience_bullets.index(para.text.strip())]

    def save_optimized_docx(self, doc: Document,  optimized_skills: List[str], output_path: str, job_title: str = "job", job_description: str = "") -> None:
        """Update the skills section in the DOCX and save it."""
        try:
            skills_started = False
            skills_index = None
            end_index = None

            # Identify where the Skills section starts and ends
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not skills_started and re.match(r"(?i)^skills\s*:", text):
                    skills_started = True
                    skills_index = i
                    continue
                if skills_started and re.match(r"^[A-Z\s]{5,}$", text):
                    end_index = i
                    break

            if skills_index is None:
                raise ValueError("Skills section not found in DOCX resume")

            if end_index is None:
                end_index = len(doc.paragraphs)

            # Delete only paragraphs after the SKILLS header and before next section
            for i in range(end_index - 1, skills_index, -1):
                p = doc.paragraphs[i]
                p._element.getparent().remove(p._element)

            # Replace the content of the skills paragraph with the first line
            skills_text_wrapped = textwrap.wrap(", ".join(optimized_skills), width=120)
            doc.paragraphs[skills_index].text = "Skills: " + skills_text_wrapped[0]

            # Add the rest of the lines after the current paragraph
            insert_after = doc.paragraphs[skills_index]._element
            for line in skills_text_wrapped[1:]:
                new_para = doc.add_paragraph(line)
                new_para.style = doc.paragraphs[skills_index].style
                insert_after.addnext(new_para._element)
                insert_after = new_para._element

            if rewrite_bullets:
                # Rewrite experience bullets to match job description
                self.rewrite_experience_bullets(doc, job_title, job_description, optimized_skills)

            if replace_job_title:
                # Replace job titles in the Professional Experience section
                self.replace_job_titles(doc, job_title)

            doc.save(output_path)
            print(f"✅ Successfully saved updated DOCX resume to: {output_path}")

        except Exception as e:
            print(f"❌ Error saving optimized DOCX resume: {e}")

    # Convert docx to pdf
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str):
        try:
            # pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
            # docx2pdf.convert(docx_path, pdf_path)
            # Use LibreOffice to convert DOCX to PDF
            import subprocess
            libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', docx_path, '--outdir', str(Path(pdf_path).parent)], check=True)
            print(f"✅ Successfully converted DOCX to PDF: {pdf_path}")
        except ImportError:
            print("❌ pypandoc module not found. Please install it using 'pip install pypandoc'.")
        except Exception as e:
            print(f"❌ Error converting DOCX to PDF: {e}")

# Function to run the resume optimization process given job description and resume path similar to main()
def run_resume_optimization(job_description: str, resume_path: str, job_id: str, job_title: str = "job", first_name: str = "first", last_name: str = "last"):
    # Get OpenAI API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("Please set the OPENAI_API_KEY environment variable")
        return

    optimizer = ResumeOptimizer(api_key)

    # Extract keywords from job description
    print("\nExtracting keywords from job description...")
    keywords = optimizer.extract_keywords(job_description)
    print(f"Top keywords identified: {', '.join(keywords)}")

    # Check if resume file exists
    if not Path(resume_path).exists():
        print("Resume file not found!")
        return

    ext = Path(resume_path).suffix.lower()
    output_path = None

    # Process resume based on file type
    print("\nProcessing resume...")
    # Sanitize variables to ensure they don't contain invalid path characters
    sanitized_job_title = re.sub(r'[\/:*?"<>|]', '_', job_title.lower().replace(' ', '_'))
    sanitized_first_name = re.sub(r'[\/:*?"<>|]', '_', first_name.lower())
    sanitized_last_name = re.sub(r'[\/:*?"<>|]', '_', last_name.lower())

    if ext == ".pdf":
        content, current_skills, format_info = optimizer.read_pdf_resume(resume_path)
        if not content or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.pdf"
        optimizer.save_optimized_pdf(resume_path, optimized_skills, format_info, str(output_path))

    elif ext == ".docx":
        doc, current_skills = optimizer.read_docx_resume(resume_path)
        if not doc or not current_skills:
            return
        optimized_skills = optimizer.optimize_skills(current_skills, keywords, job_description)
        output_path_docx = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.docx"
        output_path_pdf = Path(resume_path).parent.parent / "temp" / f"{sanitized_first_name}_{sanitized_last_name}_{sanitized_job_title}_{job_id}.pdf"
        optimizer.save_optimized_docx(doc, optimized_skills, str(output_path_docx), job_title, job_description)
        optimizer.convert_docx_to_pdf(str(output_path_docx), str(output_path_pdf))
        output_path = output_path_pdf

    else:
        print("❌ Unsupported file format. Please use PDF or DOCX.")
        return

    print(f"\n🎯 Optimized resume saved to: {output_path}")
    return output_path